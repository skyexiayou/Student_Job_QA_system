from __future__ import annotations

import re
import uuid
from pathlib import Path
from typing import Iterable, List

from .models import DocumentChunk


SUPPORTED_SUFFIXES = {".txt", ".md", ".markdown", ".pdf", ".docx", ".xlsx"}


def clean_text(text: str) -> str:
    """Normalize whitespace and remove common parser noise."""
    text = text.replace("\u3000", " ")
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
    text = re.sub(r"!\[[^\]]*]\([^)]*\)", " ", text)
    text = re.sub(r"\[([^\]]+)]\([^)]*\)", r"\1", text)
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.M)
    text = re.sub(r"^\s*(?:[-*_]\s*){3,}$", "", text, flags=re.M)
    text = re.sub(r"^\s*(?:第\s*\d+\s*页|Page\s+\d+)(?:\s*/\s*\d+)?\s*$", "", text, flags=re.I | re.M)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)


def read_text_file(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gbk"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="ignore")


def read_pdf_file(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:  # pragma: no cover - depends on optional package
        raise RuntimeError("读取 PDF 需要安装 pypdf，请执行 pip install pypdf") from exc

    reader = PdfReader(str(path))
    page_texts = []
    for index, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            page_texts.append(f"\n[第 {index} 页]\n{page_text}")
    return "\n".join(page_texts)


def read_docx_file(path: Path) -> str:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - depends on optional package
        raise RuntimeError("Reading DOCX files requires python-docx; please run pip install python-docx") from exc

    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def read_xlsx_file(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except Exception as exc:  # pragma: no cover - depends on optional package
        raise RuntimeError("Reading XLSX files requires openpyxl; please run pip install openpyxl") from exc

    workbook = load_workbook(str(path), read_only=True, data_only=True)
    rows = []
    for worksheet in workbook.worksheets:
        rows.append(f"[Sheet] {worksheet.title}")
        for row in worksheet.iter_rows(values_only=True):
            values = [str(value).strip() for value in row if value is not None and str(value).strip()]
            if values:
                rows.append(" | ".join(values))
    workbook.close()
    return "\n".join(rows)


def read_document(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".markdown"}:
        return read_text_file(path)
    if suffix == ".pdf":
        return read_pdf_file(path)
    if suffix == ".docx":
        return read_docx_file(path)
    if suffix == ".xlsx":
        return read_xlsx_file(path)
    raise ValueError(f"暂不支持的文件格式：{path.suffix}")


def split_text(text: str, chunk_size: int = 650, overlap: int = 100) -> List[str]:
    """Split by semantic separators first, then enforce a fixed-length window."""
    text = clean_text(text)
    if not text:
        return []

    separators = ["\n\n", "\n", "。", "；", "，", "、", ";", "."]
    units = [text]
    for sep in separators:
        next_units: List[str] = []
        for unit in units:
            if len(unit) <= chunk_size:
                next_units.append(unit)
                continue
            parts = [part.strip() for part in unit.split(sep) if part.strip()]
            if sep in {"。", "；", "，", "、", ";", "."}:
                parts = [part + sep for part in parts]
            next_units.extend(parts or [unit])
        units = next_units

    chunks: List[str] = []
    current = ""
    for unit in units:
        if len(current) + len(unit) + 1 <= chunk_size:
            current = f"{current}\n{unit}".strip()
        else:
            if current:
                chunks.append(current)
            while len(unit) > chunk_size:
                chunks.append(unit[:chunk_size])
                unit = unit[max(1, chunk_size - overlap) :]
            current = unit
    if current:
        chunks.append(current)

    if overlap <= 0 or len(chunks) <= 1:
        return chunks
    overlapped: List[str] = []
    previous_tail = ""
    for chunk in chunks:
        merged = f"{previous_tail}\n{chunk}".strip() if previous_tail else chunk
        overlapped.append(merged[: chunk_size + overlap])
        previous_tail = chunk[-overlap:]
    return overlapped


def build_chunks(paths: Iterable[Path], chunk_size: int, overlap: int) -> List[DocumentChunk]:
    chunks: List[DocumentChunk] = []
    for path in paths:
        if path.suffix.lower() not in SUPPORTED_SUFFIXES:
            continue
        text = read_document(path)
        title = path.stem
        for order, chunk_text in enumerate(split_text(text, chunk_size, overlap), start=1):
            chunk_id = f"{path.stem}-{order}-{uuid.uuid4().hex[:8]}"
            chunks.append(
                DocumentChunk(
                    chunk_id=chunk_id,
                    source=path.name,
                    title=title,
                    content=chunk_text,
                    metadata={"order": str(order), "path": str(path)},
                )
            )
    return chunks
